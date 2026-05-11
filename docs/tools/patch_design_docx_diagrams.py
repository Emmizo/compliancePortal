#!/usr/bin/env python3
"""Replace [mermaid]… placeholder paragraphs in Design-Document.docx with embedded PNG figures."""
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph


def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement('w:p')
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def replace_mermaid_placeholder(
    paragraph: Paragraph,
    image_path: Path,
    caption: str,
    *,
    width_inches: float = 6.2,
) -> None:
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))
    cap = insert_paragraph_after(paragraph)
    r = cap.add_run(caption)
    r.italic = True


def main() -> None:
    docs_dir = Path(__file__).resolve().parent.parent
    docx_path = docs_dir / 'Design-Document.docx'
    if not docx_path.is_file():
        raise SystemExit(f'Missing {docx_path}')

    diagrams = docs_dir / 'diagrams'
    sys_ctx = diagrams / 'system-context-architecture.png'
    workflow = diagrams / 'application-workflow-state.png'
    for p in (sys_ctx, workflow):
        if not p.is_file():
            raise SystemExit(f'Missing image: {p}')

    doc = Document(str(docx_path))
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text.startswith('[mermaid] flowchart'):
            replace_mermaid_placeholder(
                paragraph,
                sys_ctx,
                'Figure — System context. REST for ordinary requests; browsers also open STOMP over WebSocket (/ws, '
                'typically proxied together with /api) for live application status updates.',
            )
        elif text.startswith('[mermaid] stateDiagram'):
            replace_mermaid_placeholder(
                paragraph,
                workflow,
                'Figure — Application lifecycle states and transitions (roles on each edge).',
            )

    doc.save(str(docx_path))
    print(f'Updated figures in {docx_path}')


if __name__ == '__main__':
    main()
